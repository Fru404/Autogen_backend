import re
from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document


# ============================================================
# VALUE FORMATTER
# ============================================================

def format_value(value):
    """
    Convert an Excel/Python value into text suitable for Word.
    """

    if value is None:
        return ""

    try:
        if pd.isna(value):
            return ""
    except (TypeError, ValueError):
        pass

    # Pandas timestamp
    if isinstance(value, pd.Timestamp):
        return value.strftime("%d/%m/%Y")

    # Python datetime
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")

    # Numbers such as 61200.0 -> 61200
    if isinstance(value, float) and value.is_integer():
        return str(int(value))

    return str(value)


# ============================================================
# SAFE FILE NAME
# ============================================================

def safe_filename(value):
    """
    Remove characters that are not allowed in filenames.
    """

    value = str(value).strip()

    value = re.sub(
        r'[<>:"/\\|?*]',
        "_",
        value
    )

    return value


# ============================================================
# REPLACE PLACEHOLDERS
# ============================================================

def replace_placeholders_in_paragraph(
    paragraph,
    data
):
    """
    Replace {{placeholder}} in a Word paragraph.

    Handles placeholders that Word has split across
    multiple runs.
    """

    runs = paragraph.runs

    if not runs:
        return

    pattern = r"\{\{([^{}]+)\}\}"

    # --------------------------------------------------------
    # First: placeholders entirely inside individual runs
    # --------------------------------------------------------

    for run in runs:

        if not run.text:
            continue

        def replace_match(match):

            field = match.group(1).strip()

            if field in data:
                return format_value(data[field])

            return match.group(0)

        run.text = re.sub(
            pattern,
            replace_match,
            run.text
        )

    # --------------------------------------------------------
    # Rebuild paragraph text
    # --------------------------------------------------------

    runs = paragraph.runs

    full_text = "".join(
        run.text or ""
        for run in runs
    )

    if "{{" not in full_text:
        return

    matches = list(
        re.finditer(
            pattern,
            full_text
        )
    )

    if not matches:
        return

    # --------------------------------------------------------
    # Build character positions for every run
    # --------------------------------------------------------

    run_positions = []

    position = 0

    for index, run in enumerate(runs):

        text = run.text or ""

        start = position
        end = position + len(text)

        run_positions.append(
            {
                "index": index,
                "start": start,
                "end": end
            }
        )

        position = end

    # --------------------------------------------------------
    # Process right-to-left
    # --------------------------------------------------------

    for match in reversed(matches):

        field = match.group(1).strip()

        if field not in data:
            continue

        replacement = format_value(
            data[field]
        )

        placeholder_start = match.start()
        placeholder_end = match.end()

        affected = []

        for item in run_positions:

            if (
                item["end"] > placeholder_start
                and
                item["start"] < placeholder_end
            ):
                affected.append(item)

        if not affected:
            continue

        first = affected[0]
        last = affected[-1]

        first_run = runs[first["index"]]
        last_run = runs[last["index"]]

        # ----------------------------------------------------
        # Placeholder entirely within one run
        # ----------------------------------------------------

        if first["index"] == last["index"]:

            text = first_run.text

            local_start = (
                placeholder_start
                - first["start"]
            )

            local_end = (
                placeholder_end
                - first["start"]
            )

            first_run.text = (
                text[:local_start]
                + replacement
                + text[local_end:]
            )

        else:

            # ------------------------------------------------
            # First run
            # ------------------------------------------------

            first_text = first_run.text

            local_start = (
                placeholder_start
                - first["start"]
            )

            first_run.text = (
                first_text[:local_start]
                + replacement
            )

            # ------------------------------------------------
            # Middle runs
            # ------------------------------------------------

            for item in affected[1:-1]:

                runs[
                    item["index"]
                ].text = ""

            # ------------------------------------------------
            # Last run
            # ------------------------------------------------

            last_text = last_run.text

            local_end = (
                placeholder_end
                - last["start"]
            )

            last_run.text = last_text[
                local_end:
            ]


# ============================================================
# PROCESS TABLE
# ============================================================

def process_table(table, data):

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:

                replace_placeholders_in_paragraph(
                    paragraph,
                    data
                )

            for nested_table in cell.tables:

                process_table(
                    nested_table,
                    data
                )


# ============================================================
# PROCESS HEADER / FOOTER
# ============================================================

def process_header_footer(
    header_footer,
    data
):

    for paragraph in header_footer.paragraphs:

        replace_placeholders_in_paragraph(
            paragraph,
            data
        )

    for table in header_footer.tables:

        process_table(
            table,
            data
        )


# ============================================================
# PROCESS DOCUMENT
# ============================================================

def process_document(
    document,
    data
):

    # Normal paragraphs
    for paragraph in document.paragraphs:

        replace_placeholders_in_paragraph(
            paragraph,
            data
        )

    # Tables
    for table in document.tables:

        process_table(
            table,
            data
        )

    # Headers / footers
    for section in document.sections:

        process_header_footer(
            section.header,
            data
        )

        process_header_footer(
            section.footer,
            data
        )


# ============================================================
# FIND PLACEHOLDERS
# ============================================================

def get_placeholders(document):

    placeholders = set()

    pattern = r"\{\{([^{}]+)\}\}"

    def scan_paragraph(paragraph):

        text = "".join(
            run.text or ""
            for run in paragraph.runs
        )

        matches = re.findall(
            pattern,
            text
        )

        for match in matches:

            placeholders.add(
                match.strip()
            )

    def scan_table(table):

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    scan_paragraph(
                        paragraph
                    )

                for nested_table in cell.tables:

                    scan_table(
                        nested_table
                    )

    # Paragraphs
    for paragraph in document.paragraphs:

        scan_paragraph(
            paragraph
        )

    # Tables
    for table in document.tables:

        scan_table(
            table
        )

    # Headers / footers
    for section in document.sections:

        for header_footer in [
            section.header,
            section.footer
        ]:

            for paragraph in header_footer.paragraphs:

                scan_paragraph(
                    paragraph
                )

            for table in header_footer.tables:

                scan_table(
                    table
                )

    return sorted(placeholders)


# ============================================================
# READ EXCEL
# ============================================================

def read_excel(excel_file):

    df = pd.read_excel(
        excel_file
    )

    # Clean column names
    df.columns = [
        str(column).strip()
        for column in df.columns
    ]

    # Replace NaN with None
    df = df.where(
        pd.notna(df),
        None
    )

    records = []

    for row in df.to_dict(
        orient="records"
    ):

        cleaned = {}

        for key, value in row.items():

            cleaned[str(key)] = (
                format_value(value)
                if value is not None
                else ""
            )

        records.append(cleaned)

    return list(df.columns), records


# ============================================================
# GENERATE ONE DOCUMENT
# ============================================================

def generate_document(
    template_file,
    output_folder,
    data,
    row_number
):

    document = Document(
        template_file
    )

    process_document(
        document,
        data
    )

    # --------------------------------------------------------
    # Filename
    # --------------------------------------------------------

    name = format_value(
        data.get(
            "name",
            f"Client_{row_number}"
        )
    )

    account = format_value(
        data.get(
            "acc_no",
            ""
        )
    )

    name = safe_filename(
        name
    )

    account = safe_filename(
        account
    )

    if not name:

        name = f"Client_{row_number}"

    if account:

        filename = (
            f"{name}_{account}.docx"
        )

    else:

        filename = (
            f"{name}.docx"
        )

    output_path = (
        Path(output_folder)
        / filename
    )

    # --------------------------------------------------------
    # Prevent overwriting
    # --------------------------------------------------------

    counter = 1

    original_path = output_path

    while output_path.exists():

        output_path = (
            original_path.parent
            /
            f"{original_path.stem}_{counter}"
            f"{original_path.suffix}"
        )

        counter += 1

    document.save(
        str(output_path)
    )

    return output_path